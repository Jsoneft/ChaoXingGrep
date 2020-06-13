import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor  # 设置字体
nameModel = "第%s章"
targetUrl = {
    nameModel % 0: "http://exm-mayuan-ans.chaoxing.com/exam/phone/look-detail?courseId=205480312&classId=1065&examId=139796&examAnswerId=316884&protocol_v=1",
    nameModel % 1: "http://exm-mayuan-ans.chaoxing.com/exam/phone/look-detail?courseId=205480312&classId=1065&examId=139889&examAnswerId=316888&protocol_v=1",
    nameModel % 2: "http://exm-mayuan-ans.chaoxing.com/exam/phone/look-detail?courseId=205480312&classId=1065&examId=140709&examAnswerId=316891&protocol_v=1",
    nameModel % 3: "http://exm-mayuan-ans.chaoxing.com/exam/phone/look-detail?courseId=205480312&classId=1065&examId=140849&examAnswerId=316896&protocol_v=1",
    nameModel % 4: "http://exm-mayuan-ans.chaoxing.com/exam/phone/look-detail?courseId=205480312&classId=1065&examId=140402&examAnswerId=316898&protocol_v=1",
    nameModel % 5: "http://exm-mayuan-ans.chaoxing.com/exam/phone/look-detail?courseId=205480312&classId=1065&examId=141494&examAnswerId=316899&protocol_v=1",
    nameModel % 6: "http://exm-mayuan-ans.chaoxing.com/exam/phone/look-detail?courseId=205480312&classId=1065&examId=141001&examAnswerId=316904&protocol_v=1",
    # nameModel % 7: "http://exm-mayuan-ans.chaoxing.com/exam/phone/look-detail?courseId=205530827&classId=612&examId=67027&examAnswerId=177593&protocol_v=1",
    # nameModel % 9: "http://exm-mayuan-ans.chaoxing.com/exam/phone/look-detail?courseId=205531299&classId=644&examId=93276&examAnswerId=209847&protocol_v=1",
    # nameModel % 10: "http://exm-mayuan-ans.chaoxing.com/exam/phone/look-detail?courseId=205531299&classId=644&examId=93277&examAnswerId=209848&protocol_v=1",
    # nameModel % 11: "http://exm-mayuan-ans.chaoxing.com/exam/phone/look-detail?courseId=205531299&classId=644&examId=93278&examAnswerId=209849&protocol_v=1",
    # nameModel % 12: "http://exm-mayuan-ans.chaoxing.com/exam/phone/look-detail?courseId=205531299&classId=644&examId=93279&examAnswerId=209851&protocol_v=1",
    # nameModel % 13: "http://exm-mayuan-ans.chaoxing.com/exam/phone/look-detail?courseId=205531299&classId=644&examId=93280&examAnswerId=209852&protocol_v=1",
    # nameModel % 14: "http://exm-mayuan-ans.chaoxing.com/exam/phone/look-detail?courseId=205531299&classId=644&examId=93281&examAnswerId=209854&protocol_v=1",
    # nameModel % 0: "http://exm-mayuan-ans.chaoxing.com/exam/phone/look-detail?courseId=205531299&classId=644&examId=93282&examAnswerId=209855&protocol_v=1"
}

head = {
    "user-agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_5) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/83.0.4103.97 Safari/537.36",
    "Cookie": "tl=1; uname=1810311207; lv=2; fid=25368; pid=26091; _uid=94198739; uf=da0883eb5260151e367f27b2a1d2a5294ac4387bf93cc6bb41bc730575197f0e2f505b5c9976fb5398f1888311983479913b662843f1f4ad6d92e371d7fdf64402f487f7bb518a68ce915f659a7402a81471850d8bf7e34c028cbf6859df20cdc83eaf981b708130; _d=1592029700131; UID=94198739; vc=3523CA4695DE44B5DAD550E4ED9132BC; vc2=D24A3B7DD24F28CB970CAD53898933BF; vc3=HaFBMbR6JVPUMenxZ2jLvhlvPS4fT8MjJKY%2Bm%2Fae5pSY03glKtvIFjnrxUibpXaRvhTJLTZbm0E28BzjHujfmR3RlwX6ompQZ3H%2B8FnbTifY%2BeoTu8D3ZCp4ZCauAymRN7LH%2F1nZkNGdvSD9%2B8Wp2jhozlb9R%2FzYaX1VDlv%2BdYY%3Dcbd344a9073a5e5008b125fe0558dea1; xxtenc=2d5ad18f8d210b97d530175b10373df5; DSSTASH_LOG=C_38-UN_1245-US_94198739-T_1592029700132; JSESSIONID=5E5CA85CA7AE004113722EC1EAB2E0FF; source=""; route=d044092ba3e1792539d699b31c278e5b; thirdRegist=0"
}


def GetChaoXing_testBank(targetUrl_dict:dict, Conform=True):
    if Conform:
        document = Document()
    for SaveDocxName, targetUrl in targetUrl_dict.items():
        response = requests.get(targetUrl, headers=head)
        webHTML = BeautifulSoup(response.text, "html.parser")
        zm_asks = webHTML.find_all("div", attrs="zm_ask")

        qDic = {}
        count = 0
        for zm_ask in zm_asks:
            count += 1
            question = "%s、" % count
            question += zm_ask.contents[2].rstrip()
            all_ans = []
            for label in zm_ask.find_all("label"):
                line_ans = ""
                for em in label.find_all("em"):
                    line_ans += em.text
                all_ans.append(line_ans)
            qDic[question] = all_ans
        print(qDic)
        print("qDic:", len(qDic))

        zr_bgs = webHTML.find_all("div", attrs={"class": "zr_bg"})
        aList = []
        print(len(zr_bgs))
        for zr_bg in zr_bgs:
            fr = zr_bg.find("p", attrs="fr")
            answer = fr.text
            if answer.startswith("正确答案"):
                answer = answer.replace("正确答案：", "")
            else:
                answer = fr.span.text
            aList.append(answer)
        print(aList)
        print("aList:", len(aList))

        pList = []
        for key, values in qDic.items():
            pAns = []
            answers = aList.pop(0)
            find_value = False
            for answer in answers:
                for value in values:
                    if value.startswith(answer):
                        pAns.append(value)
                        find_value = True
                        break
            if not find_value:
                pAns.append(answers)
            pList.append(pAns)
        print(pList)
        print("pList:", len(pList))

        if not Conform:
            document = Document()
        document.styles['Normal'].font.name = '宋体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")

        document.add_paragraph(SaveDocxName)
        for questionStr in qDic:
            document.add_paragraph(questionStr)
            pAnsList = pList.pop(0)
            for pAns in pAnsList:
                paragraph = document.add_paragraph()
                paragraph.add_run(pAns).font.color.rgb = RGBColor(190, 0, 3)
        document.add_paragraph()
        if not Conform:
            document.save(SaveDocxName+".docx")
    if Conform:
        document.save("crawlitems/思修题库.docx")

GetChaoXing_testBank(targetUrl,Conform=True)
