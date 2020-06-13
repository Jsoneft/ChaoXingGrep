import requests
import logging
import bs4
import docx
from docx.oxml.ns import qn
from docx.shared import RGBColor

LOG_FORMAT = "%(asctime)s - %(funcName)s - %(processName)s - %(thread)s	 - %(message)s - %(msecs)d"
DATE_FORMAT = "%Y/%m/%d %H:%M:%S %p"
logging.basicConfig(level=logging.DEBUG, format=LOG_FORMAT, datefmt=DATE_FORMAT)


#  章节总数
Chapters = 8

#  目标请求网址 每次带参数都不一样, 目前是手动获取
RequestURL = {
    "第一章": "http://exm-mayuan-ans.chaoxing.com/exam/phone/look-detail?courseId=205369317&classId=613&examId=83745&examAnswerId=215680&protocol_v=1",
    "第二章": "http://exm-mayuan-ans.chaoxing.com/exam/phone/look-detail?courseId=205369317&classId=613&examId=83743&examAnswerId=217494&protocol_v=1",
    "第三章": "http://exm-mayuan-ans.chaoxing.com/exam/phone/look-detail?courseId=205369317&classId=613&examId=89164&examAnswerId=217500&protocol_v=1",
    "第四章": "http://exm-mayuan-ans.chaoxing.com/exam/phone/look-detail?courseId=205369317&classId=613&examId=89220&examAnswerId=217504&protocol_v=1",
    "第五章": "http://exm-mayuan-ans.chaoxing.com/exam/phone/look-detail?courseId=205369317&classId=613&examId=90610&examAnswerId=217509&protocol_v=1",
    "第六章": "http://exm-mayuan-ans.chaoxing.com/exam/phone/look-detail?courseId=205369317&classId=613&examId=84195&examAnswerId=217519&protocol_v=1",
    "第七章": "http://exm-mayuan-ans.chaoxing.com/exam/phone/look-detail?courseId=205369317&classId=613&examId=83744&examAnswerId=217523&protocol_v=1",
    "第八章": "http://exm-mayuan-ans.chaoxing.com/exam/phone/look-detail?courseId=205369317&classId=613&examId=84196&examAnswerId=217525&protocol_v=1",
}

#  Headers 每次相同
Headers = {
    "Host": "exm-mayuan-ans.chaoxing.com",
    "Connection": "keep-alive",
    "Upgrade-Insecure-Requests": "1",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",

    "Referer": "http://exm-mayuan-ans.chaoxing.com/selftest/look?courseId=205369317&classId=613&examAnswerId=215680"
               "&protocol_v=1",
    "Accept-Language": "zh-cn",
    "Accept-Encoding": "gzip, deflate",

    "User-Agent": "Mozilla/5.0 (iPad; CPU OS 13_5_1 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) "
                  "Mobile/15E148 ChaoXingStudy/ChaoXingStudy_3_4.5.2_ios_phone_202006052050_41 ("
                  "@Kalimdor)_14820025163915044344",

    "Cookie": "EXAM_FID=75762; jrose=23B50B91256CFE40E63D23CDDDF4BD07.self-exam-system-2057189751-s6xnf; "
              "k8s=580fee2c3873313d70f40e0abf57e4130eeec703; route=ce3aca120f3fcc9eb76807ea1ee5aae1; "
              "fanyamoocs=4DF62659665E7A4F11401F839C536D9E; DSSTASH_LOG=C_38-UN_1245-US_96643974-T_1591961140018; "
              "KI4SO_SERVER_EC=RERFSWdRQWdsckQ0aGRFcytqcUZFcjRBdlJhQi9ON2Zyc1lwYUhVaWZmblZnU25BWlRUb3dULzJP"
              "%0AUC9pQ3RCQkNPdGdxMXhwZmNPcApyTHAzald4UVJVeDl0V1AyUTVKcFhvbW5aK3Zkd2hiclJ1a0ZW"
              "%0AMy9nZ01FRXdUUWt5Q3dmUENVSXZKc0JQY3AraFFMUzhQaWVJVEh1Ck5EcWZVcTVkYUU4cVBLOCty"
              "%0AN0gwWG9jL1QrcUl5UTlkOW5PLytCa0FMNXdmTWpUU1lpMFhsb1grbG16YjVGL1VKYS9IZldLLzlv"
              "%0AMkIKeW1ZVkZ3WjJuQ2FhbjRxelFlUENzakVOdzlNZGV2Q0NnYkhxaUloeE1kYTdRVHdDRlVJQ2tw"
              "%0AUXdCZE16ZU9EVHh5RzBFVWFBMTdGcQo4QVhMYndobWRCeHRoVzljOVlxd2x4WFVyRGg0NE5QSEli"
              "%0AUVJSbGxVWlJyR0dlQ2FrUk1CMjBPQzMzdUpCRUFOSExFbndPUENzakVOCnc5TWRQMkdveGJxUHht"
              "%0AWDJGNzVlK2RHZW1ub3dVUVVJYWRoZXJJdVJhVUp6bjVwWitORTRPN3hPQldNazhQUmJsM3RydW9P"
              "%0ATW1wUTMKZWRZPT9hcHBJZD0xJmtleUlkPTE%3D; UID=96643974; _d=1591961140016; _industry=5; _tid=83397439; "
              "_uid=96643974; fid=25368; fidsCount=2; lv=2; sso_puid=96643974; "
              "uf=da0883eb5260151e367f27b2a1d2a529d1a8b41563cf1d6e6bef831c142c01789596f333c9d4772fdb14e7826cb5cd22913b662843f1f4ad6d92e371d7fdf64402f487f7bb518a68ce915f659a7402a81471850d8bf7e34c50ce84f13fd52802f2507a5220dee0ad; vc=85B166BD46F8DA9279FCC9AD4E2E41C4; vc2=7928187BC27D3C042D4311E931B63ABE; vc3=bZestdmwL6%2FdpDN84%2Bs%2BomLhXtlLweYhaaEVaK%2BkTL6qRCVrpxORUJ3jfIjURHiu%2FJiGci2q9JwUrARSUlrCP3qU%2FMvhMjWbuHGpBqTGbFXJV5RTWnYJP%2FPPluRpg5kYwFTSF8RWtzPaDmJ3kOxO3n9h%2FHBumtzbMA%2B7VYZbCoE%3D6f960c6de533834f4a7bea96d3f1ea9c; xxtenc=a4eddd8e783cc8f22837d729c82dd8a6",

}


def MainWork(document):
    for IthChapter, TarURL in RequestURL.items():
        resp = requests.get(url=TarURL, headers=Headers)
        RawHTMLText = resp.content.decode('utf-8')
        logging.debug(f"{IthChapter}'s RawHTMLText is {RawHTMLText}")
        soup = bs4.BeautifulSoup(RawHTMLText, 'html.parser')
        Questions = soup.find_all(name="div", attrs={"class": "zm_questions"})
        Answers = soup.find_all(name="p", attrs={"class": "fr"})
        assert (len(Questions) == len(Answers))
        document.add_heading(IthChapter, level=2)
        for i in range(len(Questions)):
            curq = Questions[i]
            cura = Answers[i]
            rawq = "%s." % (i+1)
            rawq += curq.contents[1].contents[2].rstrip()
            document.add_paragraph(rawq)
            CorrectOPTs = cura.text.replace("正确答案：", "").strip()
            # opts = curq.contents[1].contents[3].text.split("\n\n\n\n\n")
            opts = []
            for label in curq.find_all("label"):
                opt = ""
                for em in label.find_all("em"):
                    opt += em.text
                logging.debug(f"opt  := {opt}")
                opts.append(opt)

            out = False
            for opt in opts:
                opt = opt.strip().replace("\n", "")
                for correctOPT in CorrectOPTs:
                    if opt.startswith(correctOPT):
                        out = True
                        para = document.add_paragraph()
                        para.add_run(opt).font.color.rgb = RGBColor(190, 0, 3)

            if not out:
                para = document.add_paragraph()
                para.add_run(CorrectOPTs).font.color.rgb = RGBColor(190, 0, 3)

def CreatDOX():
    document = docx.Document()
    document.styles['Normal'].font.name = '宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
    return document


document = CreatDOX()
MainWork(document)
document.save("马原.docx")
